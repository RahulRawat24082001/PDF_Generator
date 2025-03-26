# app.py
import streamlit as st
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import io
import os
from datetime import datetime

# Custom styles for PDF
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='Header', fontSize=14, spaceAfter=12, alignment=1))
styles.add(ParagraphStyle(name='SubHeader', fontSize=12, spaceAfter=10, alignment=0))
styles.add(ParagraphStyle(name='CustomBodyText', fontSize=10, spaceAfter=8, leading=12))

# Function to generate PDF
def generate_pdf(data):
    filename = f"agreement_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
    story = []

    # Header
    story.append(Paragraph("B.K.R SUPPORT SERVICES", styles['Header']))
    story.append(Spacer(1, 12))

    # Agreement Details
    # Convert date format from "2025/03/25" to "25-03-2025"
    date_formatted = datetime.strptime(data['date_of_agreement'], '%Y/%m/%d').strftime('%d-%m-%Y')
    story.append(Paragraph(f"Date: {date_formatted}", styles['CustomBodyText']))
    story.append(Paragraph(f"Ref Number: BKR03-2025-CR702", styles['CustomBodyText']))
    story.append(Paragraph(f"Atten: {data['attention']}", styles['CustomBodyText']))
    story.append(Paragraph(f"Email: {data['email']}", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Client and Service Provider
    story.append(Paragraph(f"Client/First Party: {data['client_name']}", styles['CustomBodyText']))
    story.append(Paragraph(f"Commercial Registration Number: {data['commercial_reg_number']}", styles['CustomBodyText']))
    story.append(Paragraph(f"Second Party/Service Provider: {data['service_provider_name']}, {data['service_provider_cr']}", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Introduction
    story.append(Paragraph("Service Agreement for VAT Services and Business Support Services", styles['SubHeader']))
    story.append(Paragraph("We are thrilled to solidify our engagement with Client to provide professional services, encompassing the following terms:", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Introduction:", styles['SubHeader']))
    intro_text = """
    VAT and Business Support Services are essential for businesses looking to navigate the complexities of administrative and government-related tasks. These services, are provided by our specialized professionals, assist companies with a range of activities such as VAT filling, VAT registration and Consultancy, employee contract formalities, and compliance with local regulations. By leveraging VAT and Business support services, Your business can ensure that they remain compliant with the latest legal requirements while focusing on their core operations. This not only saves time and resources but also minimizes the risk of costly errors and delays. Whether you are a startup or an established enterprise, investing in reliable consultancy services can significantly streamline your administrative processes and contribute to your overall success.
    """
    story.append(Paragraph(intro_text, styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Scope of Services
    story.append(Paragraph("Scope of Services:", styles['SubHeader']))
    story.append(Paragraph("Our firm is committed to providing the scope of work as mentioned below in strict adherence to the terms outlined in our discussions and any subsequent correspondence. We assure you that our services will be executed with the utmost skill, care, and diligence.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Table for Scope of Work
    story.append(Paragraph("Table 1.0", styles['CustomBodyText']))
    story.append(Paragraph("Scope of work", styles['CustomBodyText']))
    story.append(Paragraph("The scope of work for VAT registration is:", styles['CustomBodyText']))
    scope_data = [
        ["1. Submission of data on the NBR portal."],
        ["2. Assistance with drafting of responses to NBR queries."],
        ["3. Review of supporting to be uploaded to the NBR"],
        ["4. Consultancy"]
    ]
    scope_table = Table(scope_data, colWidths=[6*inch])
    scope_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(scope_table)
    story.append(Spacer(1, 12))

    # VAT Registration Assessment
    story.append(Paragraph("VAT registration assessment:", styles['CustomBodyText']))
    assessment_data = [
        ["1. Detailed analysis of VAT treatment of relevant income streams."],
        ["2. Review of relevant contracts."],
        ["3. Comments on VAT legislative requirements."]
    ]
    assessment_table = Table(assessment_data, colWidths=[6*inch])
    assessment_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(assessment_table)
    story.append(Spacer(1, 12))

    # Deliverables
    story.append(Paragraph("Deliverables", styles['CustomBodyText']))
    deliverables_data = [
        ["- VAT registration"],
        ["- VAT consultancy"],
        ["- Report summarizing our findings and relevant comments."]
    ]
    deliverables_table = Table(deliverables_data, colWidths=[6*inch])
    deliverables_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(deliverables_table)
    story.append(Spacer(1, 12))

    # Additional Scope of Work Terms
    story.append(Paragraph("Any additions to the scope of work will be required through a written addendum after reaching agreement to the deliverables involved and the revised fees for services provided. Any additional services will follow the agreed terms and conditions set forth in this Service Agreement.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("The scope of work we've mutually agreed upon will be executed with unwavering professionalism, exceptional skill, meticulous attention to detail, and the requisite technical expertise.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Fee and Payment Terms
    story.append(Paragraph("Fee and Payment Terms:", styles['SubHeader']))
    story.append(Paragraph("Our professional fees are based on the degree of expertise and skills of our partners, directors and employees involved. Where a delay in provision of information causes additional time or expenses to be incurred by us in delivering the Services, we reserve the right to increase our charges to cover that additional time and expense based on our mutual agreement.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Fees Table
    fees_data = [
        ["Our office fee", "Fees"],
        ["1. VAT registration (one-time fee)", data['vat_reg_fee']],
        ["2. For the VAT registration impact assessment and Consultancy (one-time fee)", data['consultancy_fee']]
    ]
    fees_table = Table(fees_data, colWidths=[4*inch, 2*inch])
    fees_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    story.append(fees_table)
    story.append(Spacer(1, 12))

    # Billing Terms
    story.append(Paragraph("Billing terms", styles['CustomBodyText']))
    story.append(Paragraph("50% nonrefundable payment", styles['CustomBodyText']))
    story.append(Paragraph("50% before the handover of work", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Invoice will be issued on the 1st of every quarterly basis and payable within 7 days of invoice issue.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("1. VAT filling: on quarterly basis", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Additionally, any ancillary expenses incurred in the course of delivering services will be promptly reimbursed by the client. Invoices will be generated and are due for settlement upon receipt.", styles['CustomBodyText']))
    story.append(Paragraph("By engaging us to proceed with this assignment, it is acknowledged that our fee and its payment are not dependent on the outcome of our services which are Subject to ministry approvals.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Term and Termination
    story.append(Paragraph("Term and Termination:", styles['SubHeader']))
    story.append(Paragraph("Commencing on TBD, this engagement shall persist until TBD, unless terminated earlier by either party via written notice. Each party reserves the right to terminate this engagement upon [3 months] written notice to the other party.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Assumptions Regarding Scope of Work
    story.append(Paragraph("Assumptions Regarding Scope of Work", styles['SubHeader']))
    assumptions_text = """
    The laws and regulations we are providing advice on may undergo future amendments and/or be subject to different interpretations by the relevant government authorities (for instance, NBR, the Labour authority or the Ministry of Industry and Commerce). Our advice is formulated based on our interpretation of the laws, regulations, publicly available guidance, and our understanding of the prevailing practices of the relevant regulatory authority at the time of providing our advice. We cannot assure a specific outcome or anticipate all technical and interpretative challenges that the Client may encounter in the future in the event of an audit or inquiry by the relevant regulatory authorities.

    All tasks we undertake are grounded on the information furnished by the Client. Any validation or verification we conduct will be limited to sampling, and we will not authenticate all information provided to us. The services rendered under this Service Agreement do not assume a management role unless explicitly mentioned in our scope of work, and we will not serve, on a temporary or permanent basis, as a director, officer, or employee of the Client.

    The Client will bear full and sole responsibility for exercising independent business judgment regarding our services, making and executing decisions as necessary, and determining future courses of action (including regarding our recommendations) concerning any matters addressed in the deliverables submitted to the Client.

    Our work will be confined to the matters outlined in this Service Agreement. We will not be obliged to update the contents of our deliverables after their issuance date.
    In no event shall B.K.R Support Services, its partners, principals, or employees be liable for consequential, special, indirect, incidental, punitive, or exemplary damages, costs, expenses, or losses (including, without limitation, lost profits and opportunity costs).
    """
    story.append(Paragraph(assumptions_text, styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Confidentiality
    story.append(Paragraph("Confidentiality:", styles['SubHeader']))
    story.append(Paragraph("We pledge to uphold the strict confidentiality of all information shared by the client throughout our engagement, except in cases mandated by law or with explicit consent from the client.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Ownership of Work Product
    story.append(Paragraph("Ownership of Work Product:", styles['SubHeader']))
    story.append(Paragraph("Upon the full payment of all fees and expenses, all deliverables or work product generated during our engagement shall become the exclusive property of the client.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Liability
    story.append(Paragraph("Liability:", styles['SubHeader']))
    story.append(Paragraph("Our liability concerning any claim arising from or in connection with our services shall be limited to the fees paid by the client for the services giving rise to such claim. However, this limitation of liability shall not apply in the case of fraud or willful misconduct.", styles['CustomBodyText']))
    story.append(Paragraph("B.K.R Support Services (second party) will bear NO responsibility OR liability of documents or legal contracts provided by Company/Client (first party)", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Legislative Compliance
    story.append(Paragraph("Legislative Compliance:", styles['SubHeader']))
    story.append(Paragraph("Our firm acknowledges and agrees to comply with all applicable laws, regulations, and industry standards pertinent to the services rendered under this engagement. We undertake to maintain accurate records and ensure full transparency in all our dealings to mitigate any risks of legislative penalties.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Indemnification
    story.append(Paragraph("Indemnification:", styles['SubHeader']))
    story.append(Paragraph("The client agrees to indemnify and hold our firm harmless against any losses, liabilities, damages, or expenses (including reasonable attorney fees) incurred as a result of any breach of this Service Agreement or any claims arising from the client's actions or omissions.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Governing Law
    story.append(Paragraph("Governing Law:", styles['SubHeader']))
    story.append(Paragraph("This engagement shall be governed by and construed in accordance with the laws of the applicable jurisdiction. Any disputes arising from or in connection with this engagement shall be subject to the exclusive jurisdiction of the courts of Bahrain.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Distribution of Deliverables
    story.append(Paragraph("Distribution of Deliverables:", styles['SubHeader']))
    story.append(Paragraph("Our deliverables, provided in any format, are confidential and intended solely for your use. You agree not to share, reproduce, or reference them without our written consent, except for internal purposes. Sharing them doesn't grant third-party rights, and we're not liable to third parties. If you share them with a third party, both you and the recipient must sign a Hold Harmless Letter.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Timelines
    story.append(Paragraph("Timelines", styles['SubHeader']))
    story.append(Paragraph("We will mutually agree on target completion dates for the work. Our ability to meet these deadlines depends on the quality, timeliness, and availability of information. We will make every effort to adhere to agreed timetables. However, please note that timeframes provided are approximate and may be affected by the start of the engagement. Delays in receiving necessary information or access to key personnel on your end may extend the completion timeframe, for which we will not be held accountable.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Force Majeure
    story.append(Paragraph("Force Majeure:", styles['SubHeader']))
    force_majeure_text = """
    a) Either party may claim an event of force majeure. Events of force majeure are events beyond the control of either party and which either party could not foresee or reasonably provide against and which prevents either party from wholly or partly performing any duties under this Agreement, except for any events to the extent caused by intentional or gross negligent acts of the First Party, its operators, employees, or agents;
    b) The First Party claiming an event of force majeure which hinders the performance of its Services under this agreement, shall to the best of its ability give immediate written notice to the Second party of such event of force majeure including a statement describing the effect of such occurrence upon performance of this Agreement. Without prejudice to the generality of the foregoing provisions, the following events shall be recognized as events of force majeure: war, natural disasters, excluding pandemics and strikes.
    """
    story.append(Paragraph(force_majeure_text, styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Other Terms
    story.append(Paragraph("Other terms:", styles['SubHeader']))
    other_terms_data = [
        ["- All fees are Exclusive of Government Fee."],
        ["- Invoice will be generated as per service agreement."],
        ["- Services will be provided according to Main Business/One CR, any Branches are not included"],
        ["- All Jobs/requests Must be Received till 5pm, (In case any job Handed after 5 PM will be Process on second Working Day)"]
    ]
    other_terms_table = Table(other_terms_data, colWidths=[6*inch])
    other_terms_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(other_terms_table)
    story.append(Spacer(1, 12))

    # Company Bank Details
    story.append(Paragraph("Our Company Bank Details:", styles['SubHeader']))
    bank_details_data = [
        ["Account Name: B.K.R Support Services"],
        ["Bank Name: Al Salam Bank"],
        ["Account No: 294395100100"],
        ["IBAN Number: BH39ALSA00294395100100"],
        ["Swift code: ALSABHBM"],
        ["Branch: Sanabis"],
        [""],
        ["By Cheque: under name of \"B.K.R Support Services\""]
    ]
    bank_details_table = Table(bank_details_data, colWidths=[6*inch])
    bank_details_table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(bank_details_table)
    story.append(Spacer(1, 12))

    # Conclusion and Acceptance
    story.append(Paragraph("Conclusion and Acceptance:", styles['SubHeader']))
    story.append(Paragraph("Please indicate your acceptance of the terms outlined in this letter by signing and returning a copy to us. Should you require any further clarification or have additional inquiries, please do not hesitate to contact us.", styles['CustomBodyText']))
    story.append(Paragraph("We eagerly anticipate the opportunity to collaborate with you and deliver exceptional service, while ensuring full compliance with all applicable laws and regulations.", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Signature Section
    story.append(Paragraph("Yours sincerely,", styles['CustomBodyText']))
    story.append(Paragraph("On Behalf of,", styles['CustomBodyText']))
    story.append(Paragraph("B.K.R Support Services", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Director", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Sign", styles['CustomBodyText']))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Authorized Person Name: {data['authorized_person']}", styles['CustomBodyText']))
    story.append(Spacer(1, 12))

    # Footer
    footer_text = """
    p : 75069  |  t : +97333500126  |  e : infobkr@bkrgroup.co  |  w : www.bkrgroup.co
    Bldg 2196 | Office 101 | 10th Floor Road 3640 | Block 436 | Al Seef Kingdom of Bahrain
    """
    story.append(Paragraph(footer_text, styles['CustomBodyText']))

    # Build PDF
    doc.build(story)
    return filename

# Function to generate Word document
def generate_word(data):
    doc = Document()
    doc.add_heading('B.K.R SUPPORT SERVICES', 0)

    # Agreement Details
    date_formatted = datetime.strptime(data['date_of_agreement'], '%Y/%m/%d').strftime('%d-%m-%Y')
    doc.add_paragraph(f"Date: {date_formatted}")
    doc.add_paragraph(f"Ref Number: BKR03-2025-CR702")
    doc.add_paragraph(f"Atten: {data['attention']}")
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph()

    # Client and Service Provider
    doc.add_paragraph(f"Client/First Party: {data['client_name']}")
    doc.add_paragraph(f"Commercial Registration Number: {data['commercial_reg_number']}")
    doc.add_paragraph(f"Second Party/Service Provider: {data['service_provider_name']}, {data['service_provider_cr']}")
    doc.add_paragraph()

    # Introduction
    doc.add_heading("Service Agreement for VAT Services and Business Support Services", level=2)
    doc.add_paragraph("We are thrilled to solidify our engagement with Client to provide professional services, encompassing the following terms:")
    doc.add_heading("Introduction:", level=2)
    intro_text = """
    VAT and Business Support Services are essential for businesses looking to navigate the complexities of administrative and government-related tasks. These services, are provided by our specialized professionals, assist companies with a range of activities such as VAT filling, VAT registration and Consultancy, employee contract formalities, and compliance with local regulations. By leveraging VAT and Business support services, Your business can ensure that they remain compliant with the latest legal requirements while focusing on their core operations. This not only saves time and resources but also minimizes the risk of costly errors and delays. Whether you are a startup or an established enterprise, investing in reliable consultancy services can significantly streamline your administrative processes and contribute to your overall success.
    """
    doc.add_paragraph(intro_text)
    doc.add_paragraph()

    # Scope of Services
    doc.add_heading("Scope of Services:", level=2)
    doc.add_paragraph("Our firm is committed to providing the scope of work as mentioned below in strict adherence to the terms outlined in our discussions and any subsequent correspondence. We assure you that our services will be executed with the utmost skill, care, and diligence.")
    doc.add_paragraph()

    # Scope of Work
    doc.add_paragraph("Table 1.0")
    doc.add_paragraph("Scope of work")
    doc.add_paragraph("The scope of work for VAT registration is:")
    for item in [
        "1. Submission of data on the NBR portal.",
        "2. Assistance with drafting of responses to NBR queries.",
        "3. Review of supporting to be uploaded to the NBR",
        "4. Consultancy"
    ]:
        doc.add_paragraph(item)
    doc.add_paragraph()

    # VAT Registration Assessment
    doc.add_paragraph("VAT registration assessment:")
    for item in [
        "1. Detailed analysis of VAT treatment of relevant income streams.",
        "2. Review of relevant contracts.",
        "3. Comments on VAT legislative requirements."
    ]:
        doc.add_paragraph(item)
    doc.add_paragraph()

    # Deliverables
    doc.add_paragraph("Deliverables")
    for item in [
        "- VAT registration",
        "- VAT consultancy",
        "- Report summarizing our findings and relevant comments."
    ]:
        doc.add_paragraph(item)
    doc.add_paragraph()

    # Additional Scope of Work Terms
    doc.add_paragraph("Any additions to the scope of work will be required through a written addendum after reaching agreement to the deliverables involved and the revised fees for services provided. Any additional services will follow the agreed terms and conditions set forth in this Service Agreement.")
    doc.add_paragraph()
    doc.add_paragraph("The scope of work we've mutually agreed upon will be executed with unwavering professionalism, exceptional skill, meticulous attention to detail, and the requisite technical expertise.")
    doc.add_paragraph()

    # Fee and Payment Terms
    doc.add_heading("Fee and Payment Terms:", level=2)
    doc.add_paragraph("Our professional fees are based on the degree of expertise and skills of our partners, directors and employees involved. Where a delay in provision of information causes additional time or expenses to be incurred by us in delivering the Services, we reserve the right to increase our charges to cover that additional time and expense based on our mutual agreement.")
    doc.add_paragraph()

    # Fees Table (simplified for Word)
    doc.add_paragraph("Our office fee | Fees")
    doc.add_paragraph(f"1. VAT registration (one-time fee) | {data['vat_reg_fee']}")
    doc.add_paragraph(f"2. For the VAT registration impact assessment and Consultancy (one-time fee) | {data['consultancy_fee']}")
    doc.add_paragraph()

    # Billing Terms
    doc.add_paragraph("Billing terms")
    doc.add_paragraph("50% nonrefundable payment")
    doc.add_paragraph("50% before the handover of work")
    doc.add_paragraph()
    doc.add_paragraph("Invoice will be issued on the 1st of every quarterly basis and payable within 7 days of invoice issue.")
    doc.add_paragraph()
    doc.add_paragraph("1. VAT filling: on quarterly basis")
    doc.add_paragraph()
    doc.add_paragraph("Additionally, any ancillary expenses incurred in the course of delivering services will be promptly reimbursed by the client. Invoices will be generated and are due for settlement upon receipt.")
    doc.add_paragraph("By engaging us to proceed with this assignment, it is acknowledged that our fee and its payment are not dependent on the outcome of our services which are Subject to ministry approvals.")
    doc.add_paragraph()

    # Term and Termination
    doc.add_heading("Term and Termination:", level=2)
    doc.add_paragraph("Commencing on TBD, this engagement shall persist until TBD, unless terminated earlier by either party via written notice. Each party reserves the right to terminate this engagement upon [3 months] written notice to the other party.")
    doc.add_paragraph()

    # Assumptions Regarding Scope of Work
    doc.add_heading("Assumptions Regarding Scope of Work", level=2)
    assumptions_text = """
    The laws and regulations we are providing advice on may undergo future amendments and/or be subject to different interpretations by the relevant government authorities (for instance, NBR, the Labour authority or the Ministry of Industry and Commerce). Our advice is formulated based on our interpretation of the laws, regulations, publicly available guidance, and our understanding of the prevailing practices of the relevant regulatory authority at the time of providing our advice. We cannot assure a specific outcome or anticipate all technical and interpretative challenges that the Client may encounter in the future in the event of an audit or inquiry by the relevant regulatory authorities.

    All tasks we undertake are grounded on the information furnished by the Client. Any validation or verification we conduct will be limited to sampling, and we will not authenticate all information provided to us. The services rendered under this Service Agreement do not assume a management role unless explicitly mentioned in our scope of work, and we will not serve, on a temporary or permanent basis, as a director, officer, or employee of the Client.

    The Client will bear full and sole responsibility for exercising independent business judgment regarding our services, making and executing decisions as necessary, and determining future courses of action (including regarding our recommendations) concerning any matters addressed in the deliverables submitted to the Client.

    Our work will be confined to the matters outlined in this Service Agreement. We will not be obliged to update the contents of our deliverables after their issuance date.
    In no event shall B.K.R Support Services, its partners, principals, or employees be liable for consequential, special, indirect, incidental, punitive, or exemplary damages, costs, expenses, or losses (including, without limitation, lost profits and opportunity costs).
    """
    doc.add_paragraph(assumptions_text)
    doc.add_paragraph()

    # Confidentiality
    doc.add_heading("Confidentiality:", level=2)
    doc.add_paragraph("We pledge to uphold the strict confidentiality of all information shared by the client throughout our engagement, except in cases mandated by law or with explicit consent from the client.")
    doc.add_paragraph()

    # Ownership of Work Product
    doc.add_heading("Ownership of Work Product:", level=2)
    doc.add_paragraph("Upon the full payment of all fees and expenses, all deliverables or work product generated during our engagement shall become the exclusive property of the client.")
    doc.add_paragraph()

    # Liability
    doc.add_heading("Liability:", level=2)
    doc.add_paragraph("Our liability concerning any claim arising from or in connection with our services shall be limited to the fees paid by the client for the services giving rise to such claim. However, this limitation of liability shall not apply in the case of fraud or willful misconduct.")
    doc.add_paragraph("B.K.R Support Services (second party) will bear NO responsibility OR liability of documents or legal contracts provided by Company/Client (first party)")
    doc.add_paragraph()

    # Legislative Compliance
    doc.add_heading("Legislative Compliance:", level=2)
    doc.add_paragraph("Our firm acknowledges and agrees to comply with all applicable laws, regulations, and industry standards pertinent to the services rendered under this engagement. We undertake to maintain accurate records and ensure full transparency in all our dealings to mitigate any risks of legislative penalties.")
    doc.add_paragraph()

    # Indemnification
    doc.add_heading("Indemnification:", level=2)
    doc.add_paragraph("The client agrees to indemnify and hold our firm harmless against any losses, liabilities, damages, or expenses (including reasonable attorney fees) incurred as a result of any breach of this Service Agreement or any claims arising from the client's actions or omissions.")
    doc.add_paragraph()

    # Governing Law
    doc.add_heading("Governing Law:", level=2)
    doc.add_paragraph("This engagement shall be governed by and construed in accordance with the laws of the applicable jurisdiction. Any disputes arising from or in connection with this engagement shall be subject to the exclusive jurisdiction of the courts of Bahrain.")
    doc.add_paragraph()

    # Distribution of Deliverables
    doc.add_heading("Distribution of Deliverables:", level=2)
    doc.add_paragraph("Our deliverables, provided in any format, are confidential and intended solely for your use. You agree not to share, reproduce, or reference them without our written consent, except for internal purposes. Sharing them doesn't grant third-party rights, and we're not liable to third parties. If you share them with a third party, both you and the recipient must sign a Hold Harmless Letter.")
    doc.add_paragraph()

    # Timelines
    doc.add_heading("Timelines", level=2)
    doc.add_paragraph("We will mutually agree on target completion dates for the work. Our ability to meet these deadlines depends on the quality, timeliness, and availability of information. We will make every effort to adhere to agreed timetables. However, please note that timeframes provided are approximate and may be affected by the start of the engagement. Delays in receiving necessary information or access to key personnel on your end may extend the completion timeframe, for which we will not be held accountable.")
    doc.add_paragraph()

    # Force Majeure
    doc.add_heading("Force Majeure:", level=2)
    force_majeure_text = """
    a) Either party may claim an event of force majeure. Events of force majeure are events beyond the control of either party and which either party could not foresee or reasonably provide against and which prevents either party from wholly or partly performing any duties under this Agreement, except for any events to the extent caused by intentional or gross negligent acts of the First Party, its operators, employees, or agents;
    b) The First Party claiming an event of force majeure which hinders the performance of its Services under this agreement, shall to the best of its ability give immediate written notice to the Second party of such event of force majeure including a statement describing the effect of such occurrence upon performance of this Agreement. Without prejudice to the generality of the foregoing provisions, the following events shall be recognized as events of force majeure: war, natural disasters, excluding pandemics and strikes.
    """
    doc.add_paragraph(force_majeure_text)
    doc.add_paragraph()

    # Other Terms
    doc.add_heading("Other terms:", level=2)
    for item in [
        "- All fees are Exclusive of Government Fee.",
        "- Invoice will be generated as per service agreement.",
        "- Services will be provided according to Main Business/One CR, any Branches are not included",
        "- All Jobs/requests Must be Received till 5pm, (In case any job Handed after 5 PM will be Process on second Working Day)"
    ]:
        doc.add_paragraph(item)
    doc.add_paragraph()

    # Company Bank Details
    doc.add_heading("Our Company Bank Details:", level=2)
    for item in [
        "Account Name: B.K.R Support Services",
        "Bank Name: Al Salam Bank",
        "Account No: 294395100100",
        "IBAN Number: BH39ALSA00294395100100",
        "Swift code: ALSABHBM",
        "Branch: Sanabis",
        "",
        "By Cheque: under name of \"B.K.R Support Services\""
    ]:
        doc.add_paragraph(item)
    doc.add_paragraph()

    # Conclusion and Acceptance
    doc.add_heading("Conclusion and Acceptance:", level=2)
    doc.add_paragraph("Please indicate your acceptance of the terms outlined in this letter by signing and returning a copy to us. Should you require any further clarification or have additional inquiries, please do not hesitate to contact us.")
    doc.add_paragraph("We eagerly anticipate the opportunity to collaborate with you and deliver exceptional service, while ensuring full compliance with all applicable laws and regulations.")
    doc.add_paragraph()

    # Signature Section
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph("On Behalf of,")
    doc.add_paragraph("B.K.R Support Services")
    doc.add_paragraph()
    doc.add_paragraph("Director")
    doc.add_paragraph()
    doc.add_paragraph("Sign")
    doc.add_paragraph()
    doc.add_paragraph(f"Authorized Person Name: {data['authorized_person']}")
    doc.add_paragraph()

    # Footer
    footer_text = """
    p : 75069  |  t : +97333500126  |  e : infobkr@bkrgroup.co  |  w : www.bkrgroup.co
    Bldg 2196 | Office 101 | 10th Floor Road 3640 | Block 436 | Al Seef Kingdom of Bahrain
    """
    doc.add_paragraph(footer_text)

    # Save Word document
    filename = f"agreement_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

# Streamlit App
def main():
    st.set_page_config(page_title="Generator", layout="wide")
    st.markdown("""
        <style>
        .stApp {
            background-color: #1E1E1E;
            color: white;
        }
        .stTextInput > div > div > input, .stSelectbox > div > div > select {
            background-color: #2D2D2D;
            color: white;
            border: 1px solid #555;
        }
        .stButton > button {
            background-color: #4CAF50;
            color: white;
            border: none;
            padding: 10px 20px;
            text-align: center;
            text-decoration: none;
            display: inline-block;
            font-size: 16px;
            margin: 4px 2px;
            cursor: pointer;
            border-radius: 4px;
        }
        .stButton > button:hover {
            background-color: #45a049;
        }
        </style>
    """, unsafe_allow_html=True)

    st.title("Generator")

    # Form Inputs
    template = st.selectbox("Select Template", ["VAT REGISTRATION"])
    date_of_agreement = st.text_input("Date of Agreement", "2025/03/25")
    attention = st.text_input("Attention", "Nothing")
    email = st.text_input("Email", "testing@gmail.com")
    client_name = st.text_input("Client Name", "Testing Team")
    commercial_reg_number = st.text_input("Commercial Registration Number", "ABCD")
    service_provider_name = st.text_input("Service Provider Name", "None")
    service_provider_cr = st.text_input("Service Provider CR Number", "None")
    vat_reg_fee = st.text_input("VAT Registration Fee", "ASDGG")
    consultancy_fee = st.text_input("Consultancy Fee", "100")
    authorized_person = st.text_input("Authorized Person Name", "Nothing")

    # Buttons
    if st.button("Generate VAT Document"):
        data = {
            "template": template,
            "date_of_agreement": date_of_agreement,
            "attention": attention,
            "email": email,
            "client_name": client_name,
            "commercial_reg_number": commercial_reg_number,
            "service_provider_name": service_provider_name,
            "service_provider_cr": service_provider_cr,
            "vat_reg_fee": vat_reg_fee,
            "consultancy_fee": consultancy_fee,
            "authorized_person": authorized_person
        }

        # Generate PDF
        pdf_file = generate_pdf(data)
        with open(pdf_file, "rb") as f:
            st.session_state['pdf_file'] = f.read()
        os.remove(pdf_file)

        # Generate Word
        word_file = generate_word(data)
        with open(word_file, "rb") as f:
            st.session_state['word_file'] = f.read()
        os.remove(word_file)

        st.success("Document generated successfully!")

    # Download Buttons
    if 'pdf_file' in st.session_state:
        st.download_button(
            label="Download VAT Document (PDF)",
            data=st.session_state['pdf_file'],
            file_name="vat_agreement.pdf",
            mime="application/pdf"
        )

    if 'word_file' in st.session_state:
        st.download_button(
            label="Download VAT Document (Word)",
            data=st.session_state['word_file'],
            file_name="vat_agreement.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()