import streamlit as st
import os
import base64
from docx import Document
from datetime import datetime
from docx.shared import Inches
import platform
import subprocess


def replace_placeholders(doc, placeholders):
    """Replace placeholders in a Word document."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    return doc


def convert_to_pdf(doc_path, pdf_path):
    """Convert a Word document to PDF."""
    if platform.system() == "Windows":
        import comtypes.client
        import pythoncom
        pythoncom.CoInitialize()
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
    else:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf',
                       '--outdir', os.path.dirname(pdf_path), doc_path], check=True)


def main():
    st.title("PDF Document Generator")

    doc_type = st.selectbox("Select Document Type", [
                            "Invoice India", "Invoice ROW", "NDA India", "NDA ROW", "Contract India", "Contract ROW"])

    placeholders = {}
    invoice_number = int(st.session_state.get("invoice_number", 0)) + 1
    st.session_state["invoice_number"] = invoice_number

    if "Invoice India" in doc_type:
        placeholders["<<Client Name>>"] = st.text_input("Client Name")
        placeholders["<<Company Name>>"] = st.text_input("Company Name")
        placeholders["<<Invoice Number>>"] = st.text_input("Invoice Number")
        placeholders["<<Invoice Date>>"] = datetime.now().strftime("%d-%m-%Y")
        placeholders["<<Client Address>>"] = st.text_input("Client Address")
        placeholders["<<GST>>"] = st.text_input("GST Number")
        placeholders["<<Invoice Number>>"] = str(invoice_number)
        placeholders["<<Project Name>>"] = st.text_input("Project Name")
        placeholders["<<Phone Number>>"] = st.text_input("Phone Number")

    elif "Invoice ROW" in doc_type:
        placeholders["<<Client Name>>"] = st.text_input("Client Name")
        placeholders["<<Company Name>>"] = st.text_input("Company Name")
        placeholders["<<Invoice Number>>"] = st.text_input("Invoice Number")
        placeholders["<<Invoice Date>>"] = datetime.now().strftime("%d-%m-%Y")
        placeholders["<<Client Address>>"] = st.text_input("Client Address")
        placeholders["<<GST>>"] = st.text_input("GST Number")
        placeholders["<<Invoice Number>>"] = str(invoice_number)
        placeholders["<<Project Name>>"] = st.text_input("Project Name")
        placeholders["<<Phone Number>>"] = st.text_input("Phone Number")

    elif "NDA India" in doc_type:
        placeholders["<<Party A Name>>"] = st.text_input("Party A Name")
        placeholders["<<Party B Name>>"] = st.text_input("Party B Name")
        placeholders["<<Agreement Date>>"] = datetime.now().strftime(
            "%d-%m-%Y")

    elif "NDA ROW" in doc_type:
        placeholders["<<Party A Name>>"] = st.text_input("Party A Name")
        placeholders["<<Party B Name>>"] = st.text_input("Party B Name")
        placeholders["<<Agreement Date>>"] = datetime.now().strftime(
            "%d-%m-%Y")

    elif "Contract India" in doc_type:
        placeholders["<<Consultant Name>>"] = st.text_input("Consultant Name")
        placeholders["<<Company Name>>"] = st.text_input("Company Name")
        placeholders["<<Contract Date>>"] = datetime.now().strftime("%d-%m-%Y")

    elif "Contract ROW" in doc_type:
        placeholders["<<Consultant Name>>"] = st.text_input("Consultant Name")
        placeholders["<<Company Name>>"] = st.text_input("Company Name")
        placeholders["<<Contract Date>>"] = datetime.now().strftime("%d-%m-%Y")

    signature = st.file_uploader(
        "Upload Digital Signature (PNG, JPG)", type=["png", "jpg", "jpeg"])

    template_files = {
        "Invoice India": "Invoice Template - INDIA.docx",
        "Invoice ROW": "Invoice Template - ROW.docx",
        "NDA India": "NDA Template - INDIA 4.docx",
        "NDA ROW": "NDA Template - ROW 4.docx",
        "Contract India": "Contract Template - INDIA 4.docx",
        "Contract ROW": "Contract Template - ROW 4.docx"
    }

    template_path = template_files.get(doc_type)
    if st.button("Generate PDF"):
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
            doc = replace_placeholders(doc, placeholders)
            word_output = f"{doc_type}_Generated.docx"
            pdf_output = word_output.replace(".docx", ".pdf")
            doc.save(word_output)
            convert_to_pdf(word_output, pdf_output)

            with open(pdf_output, "rb") as pdf_file:
                b64_pdf = base64.b64encode(pdf_file.read()).decode('utf-8')
                href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="{pdf_output}">Download PDF</a>'
                st.markdown(href, unsafe_allow_html=True)
        else:
            st.error("Template file not found!")


if __name__ == "__main__":
    main()
